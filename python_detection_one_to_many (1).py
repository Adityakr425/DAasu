import math
import re
import string
import time
import os
from collections import defaultdict, Counter
from typing import List, Dict, Tuple, Set, Optional

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
    print("DOCX support available")
except ImportError:
    DOCX_SUPPORT = False
    print("DOCX support not available. Install: pip install python-docx")

try:
    import PyPDF2
    PDF_SUPPORT = True
    print("PDF support available")
except ImportError:
    PDF_SUPPORT = False
    print("PDF support not available. Install: pip install PyPDF2")


class TrieNode:
    def __init__(self):
        self.children = {}
        self.is_end = False
        self.frequency = 0


class Trie:
    def __init__(self):
        self.root = TrieNode()

    def insert(self, word: str) -> None:
        current = self.root
        for char in word:
            if char not in current.children:
                current.children[char] = TrieNode()
            current = current.children[char]
        current.is_end = True
        current.frequency += 1


class RabinKarpHash:
    def __init__(self, base: int = 31, mod: int = 10 ** 9 + 7):
        self.base = base
        self.mod = mod

    def compute_hash(self, text: str) -> int:
        hash_value = 0
        power = 1
        for char in text:
            hash_value = (hash_value + (ord(char) * power) % self.mod) % self.mod
            power = (power * self.base) % self.mod
        return hash_value


class WinnowingFingerprint:
    def __init__(self, k: int = 3, window_size: int = 5):
        self.k = k
        self.window_size = window_size
        self.hasher = RabinKarpHash()

    def generate_fingerprints(self, text: str) -> Set[Tuple[int, int]]:
        kgrams = self._generate_kgrams(text)
        if len(kgrams) < self.window_size:
            return set()

        hashes = [(self.hasher.compute_hash(kgram), pos) for pos, kgram in enumerate(kgrams)]

        fingerprints = set()
        for i in range(len(hashes) - self.window_size + 1):
            window = hashes[i:i + self.window_size]
            min_hash, min_pos = min(window, key=lambda x: x[0])
            fingerprints.add((min_hash, min_pos))

        return fingerprints

    def _generate_kgrams(self, text: str) -> List[str]:
        cleaned_text = re.sub(r'[^a-zA-Z0-9\s]', '', text.lower())
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()

        if len(cleaned_text) < self.k:
            return []

        return [cleaned_text[i:i + self.k] for i in range(len(cleaned_text) - self.k + 1)]


class NGramAnalyzer:
    def __init__(self, n: int = 3):
        self.n = n
        self.trie = Trie()

    def generate_ngrams(self, text: str) -> List[str]:
        words = self._tokenize(text)
        if len(words) < self.n:
            return []

        ngrams = []
        for i in range(len(words) - self.n + 1):
            ngram = ' '.join(words[i:i + self.n])
            ngrams.append(ngram)
            self.trie.insert(ngram)

        return ngrams

    def _tokenize(self, text: str) -> List[str]:
        text = text.lower()
        text = text.translate(str.maketrans('', '', string.punctuation))
        words = [word.strip() for word in text.split() if word.strip()]
        return words


class CosineSimilarity:
    def __init__(self):
        self.vocabulary = set()
        self.document_count = 0
        self.term_doc_count = defaultdict(int)

    def fit(self, documents: List[str]):
        self.vocabulary = set()
        self.document_count = len(documents)
        self.term_doc_count = defaultdict(int)

        for doc in documents:
            words = set(self._tokenize(doc))
            self.vocabulary.update(words)
            for word in words:
                self.term_doc_count[word] += 1

    def _tokenize(self, text: str) -> List[str]:
        text = text.lower()
        text = text.translate(str.maketrans('', '', string.punctuation))
        return [word for word in text.split() if word.strip()]

    def _compute_tf_idf_vector(self, document: str) -> Dict[str, float]:
        words = self._tokenize(document)
        word_count = len(words)
        tf = Counter(words)

        tf_idf_vector = {}
        for word in self.vocabulary:
            tf_score = tf[word] / word_count if word_count > 0 else 0
            idf_score = math.log(self.document_count / (self.term_doc_count[word] + 1))
            tf_idf_vector[word] = tf_score * idf_score

        return tf_idf_vector

    def similarity(self, doc1: str, doc2: str) -> float:
        vec1 = self._compute_tf_idf_vector(doc1)
        vec2 = self._compute_tf_idf_vector(doc2)

        dot_product = sum(vec1[word] * vec2[word] for word in self.vocabulary)
        magnitude1 = math.sqrt(sum(vec1[word] ** 2 for word in self.vocabulary))
        magnitude2 = math.sqrt(sum(vec2[word] ** 2 for word in self.vocabulary))

        if magnitude1 == 0 or magnitude2 == 0:
            return 0.0

        return dot_product / (magnitude1 * magnitude2)


class Document:
    def __init__(self, text: str, filename: str = ""):
        self.text = text
        self.filename = filename
        self.length = len(text)
        self.fingerprints = set()
        self.ngrams = []
        self._analyze()

    def _analyze(self):
        winnowing = WinnowingFingerprint()
        self.fingerprints = winnowing.generate_fingerprints(self.text)

        ngram_analyzer = NGramAnalyzer()
        self.ngrams = ngram_analyzer.generate_ngrams(self.text)


class SimilarityResult:
    def __init__(self):
        self.winnowing_similarity = 0.0
        self.ngram_similarity = 0.0
        self.cosine_similarity = 0.0
        self.jaccard_similarity = 0.0
        self.common_fingerprints = 0
        self.total_fingerprints = 0
        self.execution_time = 0.0
        self.average_similarity = 0.0
        self.verdict = "MINIMAL SIMILARITY"


class PlagiarismDetector:
    def __init__(self):
        self.cosine_calculator = CosineSimilarity()

    def read_docx_file(self, filepath: str) -> str:
        try:
            doc = DocxDocument(filepath)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {e}")

    def read_pdf_file(self, filepath: str) -> str:
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = []
                for page in pdf_reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text.append(extracted)
                return ' '.join(text)
        except Exception as e:
            raise Exception(f"Error reading PDF file: {e}")

    def load_document(self, filepath: str) -> Optional[Document]:
        try:
            if not os.path.exists(filepath):
                raise Exception(f"File not found: {filepath}")

            filename = os.path.basename(filepath).lower()
            if filename.endswith('.docx'):
                if not DOCX_SUPPORT:
                    raise Exception("DOCX support not available. Install python-docx")
                text = self.read_docx_file(filepath)
            elif filename.endswith('.pdf'):
                if not PDF_SUPPORT:
                    raise Exception("PDF support not available. Install PyPDF2")
                print(f"Reading PDF file: {os.path.basename(filepath)}")
                text = self.read_pdf_file(filepath)
            else:
                raise Exception("Only DOCX and PDF files are supported")

            if not text.strip():
                raise Exception("Document appears to be empty")

            doc = Document(text, os.path.basename(filepath))
            print(f"Successfully loaded {len(text)} characters, {len(text.split())} words")
            return doc
        except Exception as e:
            print(f"Error loading file {filepath}: {e}")
            return None

    def detect_plagiarism(self, doc1: Document, doc2: Document) -> SimilarityResult:
        start_time = time.time()
        result = SimilarityResult()

        result.winnowing_similarity = self._winnowing_similarity(doc1, doc2)
        result.jaccard_similarity = self._jaccard_similarity(doc1.fingerprints, doc2.fingerprints)
        result.ngram_similarity = self._ngram_similarity(doc1, doc2)
        result.cosine_similarity = self.cosine_calculator.similarity(doc1.text, doc2.text)

        result.common_fingerprints = len(doc1.fingerprints.intersection(doc2.fingerprints))
        result.total_fingerprints = len(doc1.fingerprints.union(doc2.fingerprints))
        result.execution_time = time.time() - start_time
        result.average_similarity = (result.winnowing_similarity + result.ngram_similarity) / 2
        result.verdict = self.get_verdict(result.average_similarity)

        return result

    def _winnowing_similarity(self, doc1: Document, doc2: Document) -> float:
        if not doc1.fingerprints or not doc2.fingerprints:
            return 0.0
        intersection = len(doc1.fingerprints.intersection(doc2.fingerprints))
        union = len(doc1.fingerprints.union(doc2.fingerprints))
        return intersection / union if union > 0 else 0.0

    def _jaccard_similarity(self, set1: Set, set2: Set) -> float:
        if not set1 or not set2:
            return 0.0
        intersection = len(set1.intersection(set2))
        union = len(set1.union(set2))
        return intersection / union if union > 0 else 0.0

    def _ngram_similarity(self, doc1: Document, doc2: Document) -> float:
        if not doc1.ngrams or not doc2.ngrams:
            return 0.0
        set1 = set(doc1.ngrams)
        set2 = set(doc2.ngrams)
        return self._jaccard_similarity(set1, set2)

    def get_verdict(self, avg_similarity: float) -> str:
        if avg_similarity >= 0.8:
            return "HIGH PLAGIARISM RISK"
        if avg_similarity >= 0.5:
            return "MODERATE SIMILARITY"
        if avg_similarity >= 0.2:
            return "LOW SIMILARITY"
        return "MINIMAL SIMILARITY"

    def print_similarity_report(self, result: SimilarityResult, file1: str, file2: str, doc1: Document, doc2: Document):
        print("\n" + "=" * 60)
        print(" PLAGIARISM DETECTION REPORT")
        print("=" * 60)
        print(f"Document 1: {file1}")
        print(f"Document 2: {file2}")
        print(f"Analysis Time: {result.execution_time:.4f} seconds")

        print("\n" + "-" * 40)
        print(" DOCUMENT INFORMATION")
        print("-" * 40)
        print(f"Document 1: {len(doc1.text)} characters, {len(doc1.text.split())} words")
        print(f"Document 2: {len(doc2.text)} characters, {len(doc2.text.split())} words")

        print("\n" + "-" * 40)
        print(" SIMILARITY METRICS")
        print("-" * 40)
        print(f"Winnowing Similarity: {result.winnowing_similarity:.2%}")
        print(f"N-gram Similarity: {result.ngram_similarity:.2%}")
        print(f"Cosine Similarity: {result.cosine_similarity:.2%}")
        print(f"Jaccard Similarity: {result.jaccard_similarity:.2%}")

        print("\n" + "-" * 40)
        print(" FINGERPRINT ANALYSIS")
        print("-" * 40)
        print(f"Common Fingerprints: {result.common_fingerprints}")
        print(f"Total Fingerprints: {result.total_fingerprints}")

        print("\n" + "-" * 40)
        print(" OVERALL ASSESSMENT")
        print("-" * 40)
        print(f"Average Similarity: {result.average_similarity:.2%}")
        print(f"Verdict: {result.verdict}")
        print("=" * 60)

    def get_supported_files_from_folder(self, folder_path: str, exclude_file: str = None) -> List[str]:
        if not os.path.isdir(folder_path):
            raise Exception(f"Folder not found: {folder_path}")

        supported_files = []
        exclude_abs = os.path.abspath(exclude_file) if exclude_file else None

        for name in os.listdir(folder_path):
            full_path = os.path.join(folder_path, name)
            if not os.path.isfile(full_path):
                continue

            lower_name = name.lower()
            if lower_name.endswith('.docx') or lower_name.endswith('.pdf'):
                if exclude_abs and os.path.abspath(full_path) == exclude_abs:
                    continue
                supported_files.append(full_path)

        return sorted(supported_files)

    def compare_one_to_many(self, source_file: str, target_folder: str):
        source_doc = self.load_document(source_file)
        if not source_doc:
            print("Error: Could not load source file.")
            return

        try:
            target_files = self.get_supported_files_from_folder(target_folder, exclude_file=source_file)
        except Exception as e:
            print(e)
            return

        if not target_files:
            print("No DOCX or PDF files found in the target folder.")
            return

        results = []

        print("\n" + "=" * 80)
        print(" ONE-TO-MANY PLAGIARISM CHECK")
        print("=" * 80)
        print(f"Source file: {os.path.basename(source_file)}")
        print(f"Target folder: {target_folder}")
        print(f"Files to compare: {len(target_files)}")

        for target_file in target_files:
            print(f"\nChecking against: {os.path.basename(target_file)}")
            target_doc = self.load_document(target_file)
            if not target_doc:
                print("Skipping unreadable file.")
                continue

            self.cosine_calculator.fit([source_doc.text, target_doc.text])
            result = self.detect_plagiarism(source_doc, target_doc)

            results.append((target_file, target_doc, result))

        if not results:
            print("No valid target files could be analyzed.")
            return

        results.sort(key=lambda item: item[2].average_similarity, reverse=True)

        print("\n" + "=" * 80)
        print(" SORTED RESULTS")
        print("=" * 80)
        for rank, (target_file, _, result) in enumerate(results, start=1):
            print(f"{rank}. {os.path.basename(target_file)}")
            print(f"   Average Similarity : {result.average_similarity:.2%}")
            print(f"   Winnowing          : {result.winnowing_similarity:.2%}")
            print(f"   N-gram             : {result.ngram_similarity:.2%}")
            print(f"   Cosine             : {result.cosine_similarity:.2%}")
            print(f"   Verdict            : {result.verdict}")
            print("-" * 80)

        print("\nTop detailed match report:")
        top_file, top_doc, top_result = results[0]
        self.print_similarity_report(top_result, source_file, top_file, source_doc, top_doc)


def create_sample_docx_files():
    content1 = """Data structures and algorithms are fundamental concepts in computer science. They provide the foundation for efficient problem-solving and software development. Hash tables offer O(1) average-case lookup time, making them ideal for applications requiring fast data retrieval. The Rabin-Karp algorithm uses rolling hash techniques to achieve efficient string matching with linear average-case complexity. Trees and graphs are essential data structures for representing hierarchical and networked data respectively."""

    content2 = """Data structures and algorithms form the core concepts in computer science. They provide the basis for efficient problem-solving and program development. Hash tables provide O(1) average lookup time, making them perfect for applications that require fast data access. The Rabin-Karp method uses rolling hash techniques to achieve efficient pattern matching with linear average complexity. Trees and graphs are important data structures for representing hierarchical and network data respectively."""

    content3 = """Machine learning and artificial intelligence represent cutting-edge fields in modern technology. Neural networks can learn complex patterns from data through training processes involving backpropagation and gradient descent. Deep learning models have revolutionized computer vision, natural language processing, and speech recognition applications in recent years."""

    if not DOCX_SUPPORT:
        print("Cannot create sample DOCX files because python-docx is not installed.")
        return

    try:
        doc1 = DocxDocument()
        doc1.add_heading('DSA Document 1', 0)
        doc1.add_paragraph(content1)
        doc1.save('document1.docx')

        doc2 = DocxDocument()
        doc2.add_heading('DSA Document 2', 0)
        doc2.add_paragraph(content2)
        doc2.save('document2.docx')

        doc3 = DocxDocument()
        doc3.add_heading('ML Document 3', 0)
        doc3.add_paragraph(content3)
        doc3.save('document3.docx')

        print("\nSample DOCX files created successfully!")
        print("Files created:")
        print("document1.docx - Original document about DSA")
        print("document2.docx - Similar document (paraphrased)")
        print("document3.docx - Different document about ML/AI")
        print("\nYou can now compare files in one-to-one or one-to-many mode.")
    except Exception as e:
        print(f"Error creating sample files: {e}")


def main():
    detector = PlagiarismDetector()

    print("\n" + "=" * 50)
    print(" DOCX/PDF PLAGIARISM DETECTION SYSTEM")
    print(" B.Tech 3rd Semester DSA Project")
    print("=" * 50)
    print("Algorithms: Rabin-Karp, Winnowing, N-gram Analysis, Cosine Similarity")
    print("Data Structures: Hash Tables, Trie, Sets, Arrays")
    print("File Support: DOCX and PDF files only")

    while True:
        print("\n" + "=" * 40)
        print(" MAIN MENU")
        print("=" * 40)
        print("1. Compare two files")
        print("2. Compare one file with many files")
        print("3. Create sample files")
        print("4. Exit")
        print("-" * 40)

        try:
            choice = input("Enter your choice (1-4): ").strip()

            if choice == '1':
                print("\nEnter paths of files:")
                print("Examples:")
                print("C:\\Users\\YourName\\Documents\\essay1.docx")
                print("C:\\Users\\YourName\\Desktop\\report.pdf")
                print("document1.docx")
                print()

                file1 = input("Enter path to first file: ").strip()
                file2 = input("Enter path to second file: ").strip()

                print("\nLoading and analyzing files...")
                doc1 = detector.load_document(file1)
                doc2 = detector.load_document(file2)

                if doc1 and doc2:
                    print("\nComputing similarity metrics...")
                    detector.cosine_calculator.fit([doc1.text, doc2.text])
                    result = detector.detect_plagiarism(doc1, doc2)
                    detector.print_similarity_report(result, file1, file2, doc1, doc2)
                else:
                    print("Error: Could not load one or both files!")
                    print("Please check file paths and ensure files are DOCX or PDF format.")

            elif choice == '2':
                print("\nOne-to-many mode")
                source_file = input("Enter path to source file: ").strip()
                target_folder = input("Enter folder path containing comparison files: ").strip()
                detector.compare_one_to_many(source_file, target_folder)

            elif choice == '3':
                create_sample_docx_files()

            elif choice == '4':
                print("\n" + "=" * 50)
                print("Thank you for using DOCX/PDF Plagiarism Detection System!")
                print("Project by: Code Analyzers")
                print("Course: Data Structures and Algorithms")
                print()
                print("DSA Concepts Used:")
                print("Hash Tables for O(1) fingerprint lookups")
                print("Trie for efficient n-gram storage")
                print("Rolling Hash (Rabin-Karp) for pattern matching")
                print("Winnowing algorithm for document fingerprinting")
                print("=" * 50)
                break

            else:
                print("Invalid choice! Please enter 1, 2, 3, or 4.")

        except KeyboardInterrupt:
            print("\n\nExiting...")
            break
        except Exception as e:
            print(f"An error occurred: {e}")


if __name__ == "__main__":
    main()
